# tools/word_tools.py
"""
Word document creation tools - requires python-docx
Note: Install with: pip install python-docx
"""

from pathlib import Path
from typing import Dict, Any, List, Optional
from core.base import BaseTool, ToolDefinition, ToolParameter, ToolResult, ToolResultType
from core.registry import registry

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class WordWriteTool(BaseTool):
    """Write Word documents - creates new or overwrites existing (requires python-docx)"""
    
    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(config)
    
    @property
    def definition(self) -> ToolDefinition:
        return ToolDefinition(
            name="write_word",
            description="Write a Word document (.docx) with formatted text - creates new or overwrites existing (requires python-docx library)",
            category="document_creation",
            parameters=[
                ToolParameter(
                    name="file_path",
                    description="Path to the Word file to create",
                    param_type="string",
                    required=True
                ),
                ToolParameter(
                    name="title",
                    description="Document title",
                    param_type="string",
                    required=True
                ),
                ToolParameter(
                    name="content",
                    description="Document content (supports paragraphs separated by double newlines)",
                    param_type="string",
                    required=True
                ),
                ToolParameter(
                    name="author",
                    description="Document author",
                    param_type="string",
                    required=False
                ),
                ToolParameter(
                    name="add_header",
                    description="Add header with title",
                    param_type="boolean",
                    required=False,
                    default=True
                )
            ]
        )
    
    async def execute(self, file_path: str, title: str, content: str, 
                     author: Optional[str] = None, add_header: bool = True) -> ToolResult:
        """Write Word document - creates new or overwrites existing file"""
        try:
            if not DOCX_AVAILABLE:
                return ToolResult(
                    success=False,
                    result_type=ToolResultType.ERROR,
                    content="Word document writing requires python-docx library. Install with: pip install python-docx",
                    error_message="python-docx library not available"
                )
            
            target_file = Path(file_path)
            if not target_file.suffix:
                target_file = target_file.with_suffix('.docx')
            
            # Ensure parent directory exists
            target_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Check if file exists (for informative message)
            file_existed = target_file.exists()
            
            # Create/overwrite document
            doc = Document()
            
            # Add title as heading
            if add_header:
                doc.add_heading(title, 0)
            
            # Add author if provided
            if author:
                doc.add_paragraph(f"Author: {author}")
                doc.add_paragraph()  # Empty line
            
            # Add content (split by double newlines for paragraphs)
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    # Handle simple formatting
                    para_text = para.strip()
                    
                    # Check if it's a heading (starts with #)
                    if para_text.startswith('#'):
                        # Count # symbols to determine heading level
                        level = len(para_text) - len(para_text.lstrip('#'))
                        heading_text = para_text.lstrip('#').strip()
                        doc.add_heading(heading_text, min(level, 9))
                    else:
                        # Regular paragraph
                        doc.add_paragraph(para_text)
            
            # Save document
            doc.save(str(target_file))
            
            # Create appropriate success message
            action = "overwrote" if file_existed else "created"
            
            return ToolResult(
                success=True,
                content=f"Successfully {action} Word document: {target_file}",
                result_type=ToolResultType.TEXT,
                metadata={
                    "tool": "write_word",
                    "file_path": str(target_file),
                    "title": title,
                    "author": author,
                    "paragraphs": len(paragraphs),
                    "action": action
                }
            )
            
        except Exception as e:
            return ToolResult(
                success=False,
                result_type=ToolResultType.ERROR,
                content=f"Error creating Word document: {str(e)}",
                error_message=f"Error creating Word document: {str(e)}"
            )

# Only register if python-docx is available
if DOCX_AVAILABLE:
    registry.register(WordWriteTool)
